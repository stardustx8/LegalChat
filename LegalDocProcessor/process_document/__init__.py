import azure.functions as func
from azure.search.documents import SearchClient
from azure.core.credentials import AzureKeyCredential
from azure.storage.blob import BlobServiceClient
import logging
import os
import json
import re
import requests
from docx import Document
from docx.table import Table
from PIL import Image
import base64
from io import BytesIO
from tenacity import retry, stop_after_attempt, wait_exponential
import hashlib
from typing import List, Dict, Any, Optional, Tuple

def upload_images_to_blob(images, blob_service_client, container_name, iso_code):
    """Upload images to blob storage."""
    for image in images:
        try:
            blob_name = f"images/{iso_code}/{image['filename']}"
            blob_client = blob_service_client.get_blob_client(
                container=container_name,
                blob=blob_name
            )
            
            blob_client.upload_blob(
                image['data'],
                overwrite=True,
                content_settings={'content_type': image.get('content_type', 'image/jpeg')}
            )
            
            image['blob_url'] = blob_client.url
            logging.info(f"Uploaded image {image['filename']} to blob storage")
        except Exception as e:
            logging.error(f"Failed to upload image {image['filename']}: {str(e)}")
            image['blob_url'] = ""

def split_text_into_chunks(text, max_chunk_size=2000, chunk_overlap=200):
    chunks = []
    start = 0
    while start < len(text):
        end = start + max_chunk_size
        if end < len(text):
            # Find the last space within the chunk to avoid breaking words
            while end > start and text[end] != ' ':
                end -= 1
            if end == start:  # No space found, use original end
                end = start + max_chunk_size
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        start = end - chunk_overlap if end > chunk_overlap else end
    return chunks

def extract_table_data(table: Table) -> Optional[Dict[str, Any]]:
    """Extract table data with headers, merged cells handling, and convert to Markdown and JSON."""
    try:
        if not table.rows:
            return None
        
        # Extract table data
        rows_data = []
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells]
            rows_data.append(row_cells)
        
        if not rows_data:
            return None
        
        # Detect headers (first row)
        headers = rows_data[0] if rows_data else []
        data_rows = rows_data[1:] if len(rows_data) > 1 else []
        
        # Generate table ID
        table_text = ''.join([''.join(row) for row in rows_data])
        table_id = f"table_{hashlib.md5(table_text.encode()).hexdigest()[:8]}"
        
        # Convert to Markdown
        markdown_lines = []
        if headers:
            markdown_lines.append('| ' + ' | '.join(headers) + ' |')
            markdown_lines.append('| ' + ' | '.join(['---'] * len(headers)) + ' |')
        for row in data_rows:
            # Ensure row has same number of columns as headers
            while len(row) < len(headers):
                row.append('')
            row = row[:len(headers)]  # Truncate if too many columns
            markdown_lines.append('| ' + ' | '.join(row) + ' |')
        
        markdown = '\n'.join(markdown_lines)
        
        # Convert to JSON
        json_data = []
        for row in data_rows:
            row_dict = {}
            for i, header in enumerate(headers):
                if i < len(row):
                    row_dict[header] = row[i]
                else:
                    row_dict[header] = ''
            json_data.append(row_dict)
        
        return {
            'id': table_id,
            'headers': headers,
            'markdown': markdown,
            'json': json_data
        }
    except Exception as e:
        logging.error(f"Error extracting table data: {str(e)}")
        return None

def extract_images_from_docx(doc: Document, enable_captioning: bool) -> List[Dict[str, Any]]:
    """Extract images from DOCX document."""
    images = []
    try:
        # Extract images from document relationships
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image_part = rel.target_part
                image_data = image_part.blob
                
                # Generate image ID
                img_hash = hashlib.md5(image_data).hexdigest()[:8]
                img_id = f"figure_{img_hash}"
                filename = f"image_{img_hash}.{image_part.content_type.split('/')[-1]}"
                
                images.append({
                    'id': img_id,
                    'filename': filename,
                    'data': image_data,
                    'content_type': image_part.content_type
                })
    except Exception as e:
        logging.error(f"Error extracting images: {str(e)}")
    
    return images

@retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
def generate_image_captions(images, openai_endpoint, openai_key, deployment_name):
    # Chat completions (vision) endpoint
    chat_api_version = os.environ.get("OPENAI_API_VERSION", "2024-02-15-preview")
    openai_chat_url = (
        f"{openai_endpoint}/openai/deployments/{deployment_name}/chat/completions?api-version={chat_api_version}"
        if deployment_name else None
    )
    
    if not openai_chat_url:
        return images
    
    for image in images:
        try:
            b64 = base64.b64encode(image['data']).decode("utf-8")
            prompt_user_text = (
                "Provide a concise caption for this legal document image, and extract any "
                "legible text exactly as OCR. Return JSON with keys 'caption' and 'image_text'."
            )
            payload_chat = {
                "messages": [
                    {"role": "system", "content": "You caption legal document images and extract exact text."},
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": prompt_user_text},
                            {"type": "image_url", "image_url": {"url": f"data:{image['content_type']};base64,{b64}"}}
                        ]
                    }
                ],
                "temperature": 0.0,
                "max_tokens": 1500,
                "response_format": {"type": "json_object"}
            }
            cr = requests.post(openai_chat_url, headers={"Content-Type": "application/json", "api-key": openai_key}, json=payload_chat, timeout=45)
            cr.raise_for_status()
            content = cr.json()["choices"][0]["message"]["content"].strip()
            try:
                obj = json.loads(content)
                caption = (obj.get("caption") or "").strip()
                image_text = (obj.get("image_text") or "").strip()
            except Exception:
                caption, image_text = content, ""
            image['caption'] = caption
            image['ocr_text'] = image_text
        except Exception as ce:
            logging.warning(f"Caption/OCR failed for {image['filename']}: {ce}")
    
    return images

def extract_document_elements(blob_content, doc_name, storage_connection_string, enable_captioning):
    """Extract text blocks, tables, and images from a DOCX file using python-docx."""
    try:
        elements = []
        with BytesIO(blob_content) as doc_io:
            doc = Document(doc_io)
            
            # Process document elements in order
            for element in doc.element.body:
                if element.tag.endswith('p'):
                    # Paragraph
                    para = None
                    for p in doc.paragraphs:
                        if p._element == element:
                            para = p
                            break
                    if para and para.text.strip():
                        elements.append({
                            'type': 'text',
                            'content': para.text,
                            'metadata': {}
                        })
                
                elif element.tag.endswith('tbl'):
                    # Table
                    table = None
                    for t in doc.tables:
                        if t._element == element:
                            table = t
                            break
                    if table:
                        table_data = extract_table_data(table)
                        if table_data:
                            elements.append({
                                'type': 'table',
                                'content': table_data['markdown'],
                                'metadata': {
                                    'table_id': table_data['id'],
                                    'headers': table_data['headers'],
                                    'json_data': table_data['json']
                                }
                            })
            
            # Extract images and add as elements with data attached
            images = extract_images_from_docx(doc, enable_captioning)
            for img in images:
                elements.append({
                    'type': 'image',
                    'content': f"Image: {img['filename']}",
                    'metadata': {
                        'figure_id': img['id'],
                        'filename': img['filename'],
                        'ocr_text': ''
                    },
                    'data': img['data'],  # Attach image data
                    'content_type': img['content_type']
                })
            
            return elements
    except Exception as e:
        logging.error(f"Error extracting document elements: {str(e)}")
        return []

def main(myblob: func.InputStream):
    logging.info(f"Python blob trigger function processed blob")
    logging.info(f"Name: {myblob.name}")
    logging.info(f"Size: {myblob.length} Bytes")

    # Extract ISO code from filename
    filename = myblob.name.split('/')[-1]
    match = re.match(r"([A-Z]{2})\.docx", filename)
    if not match:
        logging.error(f"Invalid filename format: {filename}. Expected 'XX.docx' where XX is a 2-letter ISO code.")
        return
    
    iso_code = match.group(1)
    logging.info(f"Processing document for ISO code: {iso_code}")
    
    # Caption/OCR is always enabled
    enable_captioning = True
    logging.info(f"Caption/OCR is enabled by default")

    # Azure Cognitive Search and OpenAI settings from environment variables
    search_endpoint = os.environ.get("KNIFE_SEARCH_ENDPOINT")
    search_key = os.environ.get("KNIFE_SEARCH_KEY")
    search_index_name = os.environ.get("KNIFE_SEARCH_INDEX")
    openai_endpoint = os.environ.get("KNIFE_OPENAI_ENDPOINT")
    openai_key = os.environ.get("KNIFE_OPENAI_KEY")
    openai_embedding_deployment = os.environ.get("KNIFE_OPENAI_DEPLOY")
    openai_chat_deployment = os.environ.get("OPENAI_CHAT_DEPLOY")
    storage_connection_string = os.environ.get("KNIFE_STORAGE_CONNECTION_STRING")

    # Check environment variables and log missing
    env_vars = {
        "KNIFE_SEARCH_ENDPOINT": search_endpoint,
        "KNIFE_SEARCH_KEY": search_key,
        "KNIFE_SEARCH_INDEX": search_index_name,
        "KNIFE_OPENAI_ENDPOINT": openai_endpoint,
        "KNIFE_OPENAI_KEY": openai_key,
        "KNIFE_OPENAI_DEPLOY": openai_embedding_deployment
    }
    
    missing_vars = [key for key, value in env_vars.items() if not value]
    if missing_vars:
        logging.error(f"Missing required environment variables: {', '.join(missing_vars)}")
        return

    # Prepare OpenAI REST API headers and endpoint
    openai_headers = {
        "Content-Type": "application/json",
        "api-key": openai_key
    }
    openai_url = f"{openai_endpoint}/openai/deployments/{openai_embedding_deployment}/embeddings?api-version=2023-05-15"
    
    # Initialize Search client
    search_credential = AzureKeyCredential(search_key)
    search_client = SearchClient(endpoint=search_endpoint, index_name=search_index_name, credential=search_credential)

    try:
        # Extract document elements (text, tables, images)
        elements = extract_document_elements(
            myblob.read(), 
            myblob.name, 
            storage_connection_string,
            enable_captioning and openai_chat_deployment is not None
        )
        
        if not elements:
            logging.warning(f"No content extracted from {myblob.name}")
            return
        
        # Separate images from other elements
        image_elements = [e for e in elements if e['type'] == 'image']
        content_elements = [e for e in elements if e['type'] != 'image']
        
        # Process images: generate captions and upload to blob
        if openai_chat_deployment and image_elements:
            logging.info(f"Processing {len(image_elements)} images")
            
            # Prepare image data for caption generation
            image_data = []
            for elem in image_elements:
                image_data.append({
                    'id': elem['metadata']['figure_id'],
                    'filename': elem['metadata']['filename'],
                    'data': elem.get('data', b''),
                    'content_type': elem.get('content_type', 'image/jpeg')
                })
            
            # Generate captions and OCR
            image_data = generate_image_captions(
                image_data,
                openai_endpoint,
                openai_key,
                openai_chat_deployment
            )
            
            # Upload images to blob storage
            blob_service_client = BlobServiceClient.from_connection_string(storage_connection_string)
            upload_images_to_blob(image_data, blob_service_client, "legaldocsrag", iso_code)
            
            # Update image elements with captions and URLs
            for elem in image_elements:
                for img in image_data:
                    if elem['metadata'].get('figure_id') == img['id']:
                        elem['content'] = img.get('caption', elem['content'])
                        caption = img.get('caption', '')
                        ocr_text = img.get('ocr_text', '')
                        if caption and ocr_text:
                            elem['content'] = f"{caption}\n\n{ocr_text}"
                        elif caption:
                            elem['content'] = caption
                        elif ocr_text:
                            elem['content'] = ocr_text
                        else:
                            # retain existing placeholder like "Image: filename" if both are empty
                            elem['content'] = elem['content']
                        elem['metadata']['ocr_text'] = img.get('ocr_text', '')
                        elem['metadata']['blob_url'] = img.get('blob_url', '')

        # Create chunks from elements
        chunks = []
        current_chunk = []
        current_size = 0
        max_chunk_size = 2000
        
        for elem in content_elements:
            elem_text = elem['content']
            elem_size = len(elem_text)
            
            if elem['type'] == 'table':
                # Tables go in their own chunk
                if current_chunk:
                    chunks.append({
                        'text': '\n\n'.join([e['content'] for e in current_chunk]),
                        'metadata': {'chunk_type': 'text'}
                    })
                    current_chunk = []
                    current_size = 0
            
                chunks.append({
                    'text': elem_text,
                    'metadata': {
                        'chunk_type': 'table',
                        'table_id': elem['metadata'].get('table_id', ''),
                        'table_json': json.dumps(elem['metadata'].get('json_data', []))
                    }
                })
            elif current_size + elem_size > max_chunk_size and current_chunk:
                # Create a new chunk
                chunks.append({
                    'text': '\n\n'.join([e['content'] for e in current_chunk]),
                    'metadata': {'chunk_type': 'text'}
                })
                current_chunk = [elem]
                current_size = elem_size
            else:
                current_chunk.append(elem)
                current_size += elem_size
        
        # Add remaining chunk
        if current_chunk:
            chunks.append({
                'text': '\n\n'.join([e['content'] for e in current_chunk]),
                'metadata': {'chunk_type': 'text'}
            })
        
        # Add image chunks
        for elem in image_elements:
            chunks.append({
                'text': elem['content'],
                'metadata': {
                    'chunk_type': 'image',
                    'figure_id': elem['metadata'].get('figure_id', ''),
                    'ocr_text': elem['metadata'].get('ocr_text', ''),
                    'blob_url': elem['metadata'].get('blob_url', '')
                }
            })
        
        # Generate embeddings for each chunk
        embeddings = []
        for chunk_data in chunks:
            chunk_text = chunk_data['text']
            # Make direct REST API call to Azure OpenAI
            payload = {
                "input": chunk_text
            }
            response = requests.post(openai_url, headers=openai_headers, json=payload)
            response.raise_for_status()
            embedding_data = response.json()
            embedding = embedding_data['data'][0]['embedding']
            embeddings.append(embedding)
        
        # Delete existing documents for this ISO code
        logging.info(f"Deleting existing documents for ISO code: {iso_code}")
        filter_expr = f"iso_code eq '{iso_code}'"
        results = search_client.search(search_text="*", filter=filter_expr, select=["id"])
        docs_to_delete = [doc["id"] for doc in results]
        
        if docs_to_delete:
            logging.info(f"Deleting {len(docs_to_delete)} existing documents")
            delete_result = search_client.delete_documents(documents=[{"id": doc_id} for doc_id in docs_to_delete])
            logging.info(f"Delete result: {delete_result}")
        
        # Prepare documents for indexing
        documents = []
        for i, chunk_data in enumerate(chunks):
            doc = {
                "id": f"{iso_code}_{i}",
                "iso_code": iso_code,
                "chunk": chunk_data['text'],  
                "embedding": embeddings[i] if i < len(embeddings) else [0] * 3072,  
                "chunk_type": chunk_data['metadata'].get('chunk_type', 'text')
            }
            
            # Add table markdown if it's a table
            if chunk_data['metadata'].get('chunk_type') == 'table':
                doc['table_md'] = chunk_data['text']  

            documents.append(doc)

        # Upload new documents
        if documents:
            logging.info(f"Uploading {len(documents)} new documents to the search index")
            upload_result = search_client.upload_documents(documents=documents)
            success_count = sum(1 for r in upload_result if r.succeeded)
            logging.info(f"Successfully uploaded {success_count}/{len(documents)} documents")
            
            if success_count < len(documents):
                failed = [r for r in upload_result if not r.succeeded]
                logging.error(f"Failed uploads: {failed}")
        
        logging.info(f"Document processing completed for {filename}")
        
    except Exception as e:
        logging.error(f"Error processing document: {str(e)}")
        import traceback
        logging.error(f"Traceback: {traceback.format_exc()}")
        raise